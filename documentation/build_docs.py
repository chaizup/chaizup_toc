# -*- coding: utf-8 -*-
"""
Generate a comprehensive Word (.docx) feature reference for the Chaizup TOC app.
Covers every macro feature (pages, reports, engines, schedulers, doctypes) and
micro feature (formulas, validations, custom fields, hooks) with purpose,
calculation/logic, and trigger time. All terms in full form; the short form
is given in brackets the first time it appears.

Run with the bench virtualenv python:
    /workspace/development/frappe-bench/env/bin/python build_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ---------------------------------------------------------------------------
# Palette (the "Operator Cockpit" slate + indigo aesthetic)
# ---------------------------------------------------------------------------
INDIGO = RGBColor(0x4F, 0x46, 0xE5)
SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)
SLATE_600 = RGBColor(0x47, 0x55, 0x69)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "4F46E5"
SUBHEAD_FILL = "E0E7FF"
ZEBRA_FILL = "F1F5F9"
CODE_FILL = "F8FAFC"

doc = Document()

# Base styles
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = SLATE_900
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.12

for lvl, size in ((1, 20), (2, 15), (3, 12.5)):
    st = doc.styles["Heading %d" % lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = INDIGO if lvl == 1 else SLATE_900


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, size=9.5, fill=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    runs = text.split("\n")
    for i, line in enumerate(runs):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        if color:
            r.font.color.rgb = color
    if fill:
        _shade(cell, fill)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def para(text, bold=False, italic=False, color=None, size=10.5, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.12)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE_900
    # light fill via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_FILL)
    pPr.append(shd)
    return p


def table(headers, rows, widths=None, header_fill=HEADER_FILL, zebra=True, font=9.0):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        _set_cell_text(hdr[i], htext, bold=True, color=WHITE, size=font + 0.5, fill=header_fill)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            fill = ZEBRA_FILL if (zebra and ri % 2 == 1) else None
            _set_cell_text(cells[ci], str(val), size=font, fill=fill)
    if widths:
        for row in t.rows:
            for ci, w in enumerate(widths):
                row.cells[ci].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def spacer(pts=4):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


# ===========================================================================
# COVER
# ===========================================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(110)
r = t.add_run("Chaizup TOC")
r.font.size = Pt(40)
r.bold = True
r.font.color.rgb = INDIGO

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Theory of Constraints (TOC) Buffer Management for ERPNext")
r.font.size = Pt(15)
r.font.color.rgb = SLATE_600

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Complete Feature, Calculation, Logic & Trigger Reference")
r.font.size = Pt(13)
r.bold = True
r.font.color.rgb = SLATE_900

sub3 = doc.add_paragraph()
sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub3.add_run(
    "Every macro feature (pages, reports, automation engines, scheduled jobs) and\n"
    "every micro feature (formulas, validations, custom fields, hooks), with purpose,\n"
    "calculation, logic and trigger time. Abbreviations are written in full; the short\n"
    "form is given in brackets the first time it appears."
)
r.font.size = Pt(10.5)
r.italic = True
r.font.color.rgb = SLATE_500

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_before = Pt(90)
r = meta.add_run("Generated %s  ·  Application version 1.0.0  ·  Requires Frappe + ERPNext"
                 % datetime.date.today().isoformat())
r.font.size = Pt(9.5)
r.font.color.rgb = SLATE_500

doc.add_page_break()

# ===========================================================================
# 1. WHAT THE APP IS
# ===========================================================================
h1("1. What This Application Is and Why It Exists")

para(
    "Chaizup TOC is a Theory of Constraints (TOC) add-on for ERPNext, built for "
    "food and fast-moving consumer goods (FMCG) manufacturing. It replaces ERPNext's "
    "default reorder-level based automatic Material Request (MR) creation with a "
    "demand-driven replenishment engine that keeps inventory inside dynamic buffers "
    "instead of fixed minimum/maximum reorder points."
)
para("The application automates four jobs that the standard manufacturing flow does poorly:", bold=True)
bullet("keep inventory in dynamic buffers that resize themselves to real consumption, instead of static reorder points.", "Demand-driven Material Requests — ")
bullet("turn submitted Sales Projections (forecasts) and pending Sales Orders (SO) into Draft Production Plans (PP) automatically, per item and per warehouse.", "Production planning at item × warehouse level — ")
bullet("Sales Order, Work Order (WO) and Purchase Order (PO) “what counts as still pending” status lists are configured once in TOC Settings and shared by every report and every scheduled job.", "One source of truth for “pending” — ")
bullet("every automation run writes one audit log row (with per-item child rows) so any missing or unexpected document can be traced back.", "Audit-grade run logs — ")

para(
    "Important boundary: this application only adds behaviour on top of ERPNext. It does "
    "not contain a separate inventory database. Every quantity it shows is read live from "
    "ERPNext tables (Bin stock, Work Orders, Sales Orders, Purchase Orders, Stock Ledger "
    "Entries, Bills of Materials) at the moment you ask for it.", italic=True
)

# ===========================================================================
# 2. GLOSSARY
# ===========================================================================
h1("2. Glossary — Every Abbreviation in Full")
para(
    "This table is the master key. Throughout the rest of the document the full form is "
    "used, with the short form in brackets the first time it appears in a section."
)
glossary = [
    ("Theory of Constraints (TOC)", "A management method that focuses planning on the single resource that limits output (the constraint), and protects it with inventory and time buffers."),
    ("Average Daily Usage (ADU)", "How many units of an item are consumed or shipped per day, averaged over a chosen look-back window."),
    ("Replenishment Lead Time (RLT)", "Number of days from raising a replenishment signal to the goods being available in the warehouse."),
    ("Variability Factor (VF)", "A safety multiplier (typically 1.0 to 2.0) that enlarges the buffer to absorb demand and supply variability."),
    ("Demand Adjustment Factor (DAF)", "A seasonal multiplier on the buffer (for example 1.6 for a festival season) to plan for predictable demand swings."),
    ("Target Buffer", "The planned amount of stock to hold for an item at a warehouse."),
    ("Inventory Position (IP)", "The true stock picture: on-hand plus incoming supply minus committed demand."),
    ("Buffer Penetration Percentage (BP%)", "How deeply current stock has fallen into the target buffer. Higher means more urgent."),
    ("Stock Remaining Percentage (SR%)", "The inverse of Buffer Penetration; how much of the buffer is still covered (100 minus Buffer Penetration)."),
    ("Throughput per Constraint Unit (T/CU)", "Profit earned per minute of constraint-machine time; used as a tie-breaker when two items are equally urgent."),
    ("Totally Variable Cost (TVC)", "Only the costs that change per extra unit produced (raw material and packaging), not labour, rent or electricity."),
    ("Dynamic Buffer Management (DBM)", "Weekly automatic resizing of the target buffer based on recent zone behaviour."),
    ("Too Much Red (TMR)", "A DBM trigger: the buffer was in the urgent zone too often, so it is increased."),
    ("Too Much Green (TMG)", "A DBM trigger: the buffer was comfortable for a long stretch, so it is decreased."),
    ("Minimum Order Quantity (MOQ)", "The smallest batch the system is allowed to order or produce for an item at a warehouse."),
    ("Material Request (MR)", "An ERPNext document requesting goods, of type Purchase or Manufacture."),
    ("Production Plan (PP)", "An ERPNext document that groups demand and spawns Work Orders."),
    ("Work Order (WO)", "An ERPNext document for one manufacturing run."),
    ("Purchase Order (PO)", "An ERPNext document ordering goods from a supplier."),
    ("Sales Order (SO)", "An ERPNext document recording a confirmed customer order."),
    ("Delivery Note (DN)", "An ERPNext document recording goods actually shipped to a customer."),
    ("Stock Entry (SE)", "An ERPNext document recording a stock movement (issue, receipt, transfer, manufacture)."),
    ("Stock Ledger Entry (SLE)", "The atomic record ERPNext writes for every change in stock."),
    ("Bill of Materials (BOM)", "The recipe: which components and how much go into one unit of a manufactured item."),
    ("Finished Goods (FG)", "Sellable end product."),
    ("Semi-Finished Goods (SFG)", "An intermediate product (for example a blend) consumed inside a finished good."),
    ("Raw Material (RM)", "A purchased input consumed in production."),
    ("Packaging Material (PM)", "Pouches, cartons and similar purchased packaging inputs."),
    ("Unit of Measurement (UOM)", "The unit a quantity is expressed in (Gram, Kilogram, Pieces, Carton, etc.)."),
    ("Work In Process (WIP)", "Items currently inside open Work Orders, not yet finished."),
    ("Work Order Kitting Planner (WKP)", "The page that simulates whether open Work Orders can be fully kitted from available materials."),
    ("Production Overview (POR)", "The single-screen production planning console page."),
]
table(["Term (short form)", "Meaning"], glossary, widths=[2.4, 4.6], font=9.0)

# ===========================================================================
# 3. ARCHITECTURE
# ===========================================================================
h1("3. Architecture and Package Layout")
para(
    "The application is a standard Frappe app. Business logic is pure Python in the "
    "engine layer; the user interface is a mix of Frappe Desk Pages and Script Reports; "
    "all wiring to ERPNext happens through hooks."
)
table(
    ["Folder", "Responsibility"],
    [
        ("hooks.py", "Master wiring — scheduler jobs, document-event hooks, the reorder override, client scripts, fixtures."),
        ("toc_engine/", "Pure calculation engine: buffer maths, Dynamic Buffer Management, Material Request generation, component-shortage Material Requests, the Production Plan / shortage automation."),
        ("tasks/", "The scheduled background jobs (daily and weekly)."),
        ("overrides/", "Behaviour added to existing ERPNext doctypes (Item, Material Request, Production Plan, Purchase Order, the reorder routine)."),
        ("api/", "Whitelisted endpoints called by the pages and by external clients."),
        ("chaizup_toc/doctype/", "Custom doctypes (TOC Settings, TOC Item Buffer, TOC Buffer Log, Sales Projection and child tables, Item Minimum Order Quantity)."),
        ("chaizup_toc/page/", "Interactive Desk Pages (dashboards, planners, trackers)."),
        ("chaizup_toc/report/", "Five Script Reports."),
        ("setup/", "Install and uninstall routines: custom fields, roles, number cards, dashboard charts."),
        ("fixtures/", "Custom Fields, Property Setters and List View Settings shipped with the app."),
        ("patches/", "One-shot database migrations."),
        ("public/", "Client-side assets (global branding script, per-form scripts, styles, logo)."),
    ],
    widths=[2.0, 5.0],
)
para("Application identity (from hooks.py):", bold=True)
code(
    'app_name        = "chaizup_toc"\n'
    'app_title       = "Chaizup TOC"\n'
    'app_publisher   = "Chaizup"\n'
    'app_description = "Theory of Constraints Buffer Management for ERPNext"\n'
    'app_version     = "1.0.0"\n'
    'required_apps   = ["frappe", "erpnext"]'
)

# ===========================================================================
# 4. THE FORMULA ENGINE
# ===========================================================================
h1("4. The Calculation Formulas (F1 to F8)")
para(
    "These eight formulas are the mathematical heart of the application. Every quantity is "
    "computed in the item's stock Unit of Measurement (UOM) unless stated otherwise. The "
    "engine is universal: it always reads every supply and demand source for every item, "
    "regardless of whether the item is normally purchased or manufactured — unused sources "
    "simply return zero."
)

h2("4.1 Formula summary")
table(
    ["Identifier", "Full name", "Formula (plain words)", "Where it runs"],
    [
        ("F1", "Target Buffer", "Average Daily Usage × Replenishment Lead Time × Variability Factor", "buffer_calculator._target_buffer()"),
        ("F2", "Inventory Position", "On-Hand + Work In Process + On-Order − Backorders − Committed", "buffer_calculator.get_inventory_position()"),
        ("F3", "Buffer Penetration Percentage", "(Target Buffer − Inventory Position) ÷ Target Buffer × 100", "_calculate_single()"),
        ("F3a", "Stock Remaining Percentage", "Inventory Position ÷ Target Buffer × 100  (= 100 − Buffer Penetration)", "_calculate_single()"),
        ("F4", "Order Quantity", "maximum of (Target Buffer − Inventory Position) and 0", "_calculate_single() / _create_mr()"),
        ("F5", "Throughput per Constraint Unit", "(Selling Price − Totally Variable Cost) × Constraint Speed", "overrides/item.on_item_validate()"),
        ("F6", "Demand-Adjusted Buffer", "Target Buffer × Demand Adjustment Factor", "TOCItemBuffer.calculate_adjusted_buffer()"),
        ("F7", "Dynamic Buffer up-shift (Too Much Red)", "new Target = round(Target × 1.33)  [+33%]", "dbm_engine._evaluate_single()"),
        ("F8", "Dynamic Buffer down-shift (Too Much Green)", "new Target = max(floor, round(Target × 0.67))  [−33%]", "dbm_engine._evaluate_single()"),
    ],
    widths=[0.7, 1.9, 2.9, 1.5],
)

h2("4.2 Inventory Position (F2) — the five components")
para("Inventory Position is the honest stock picture. It is built from five quantities, all read live from ERPNext:")
table(
    ["Component", "Meaning", "Read from"],
    [
        ("On-Hand", "Physical stock right now.", "Bin.actual_qty (warehouses classified as Inventory)"),
        ("Work In Process (WIP)", "Output still expected from open Work Orders.", "Work Order.qty − produced_qty, plus stock in Work-In-Process warehouses"),
        ("On-Order", "Goods on open Purchase Orders.", "Bin.ordered_qty"),
        ("Backorders", "Customer demand reserved against stock.", "Bin.reserved_qty"),
        ("Committed", "Components reserved by open Work Orders.", "Work Order Item.required_qty − transferred_qty"),
    ],
    widths=[1.5, 2.7, 2.8],
)

h2("4.3 Zone classification (from Buffer Penetration)")
para(
    "Once Buffer Penetration Percentage (F3) is known, the item is placed in a colour zone. "
    "The two thresholds (default 67 and 33) live in TOC Settings and are adjustable."
)
table(
    ["Zone", "Condition (default thresholds)", "Meaning / action"],
    [
        ("Green", "Buffer Penetration below 33%", "Comfortable — no action."),
        ("Yellow", "33% to below 67%", "Plan replenishment."),
        ("Red", "67% to below 100%", "Produce or order now."),
        ("Black", "100% or more", "Stockout — Inventory Position is zero or negative. Emergency."),
    ],
    widths=[1.0, 2.3, 3.7],
)

h2("4.4 A full worked example")
para("Finished Goods (FG) item, manufactured, at the Finished Goods store, during a festival season:")
code(
    "Inputs:\n"
    "  Average Daily Usage (ADU) = 10 units/day\n"
    "  Replenishment Lead Time (RLT) = 7 days\n"
    "  Variability Factor (VF) = 1.5\n"
    "  Demand Adjustment Factor (DAF) = 1.6  (festival uplift)\n\n"
    "F1  Target Buffer        = 10 × 7 × 1.5      = 105 units\n"
    "F6  Demand-Adjusted      = 105 × 1.6         = 168 units  (effective buffer)\n\n"
    "Live stock snapshot:\n"
    "  On-Hand 42, Work In Process 18, On-Order 0, Backorders 8, Committed 5\n"
    "F2  Inventory Position   = 42 + 18 + 0 − 8 − 5 = 47 units\n\n"
    "F3  Buffer Penetration   = (168 − 47) / 168 × 100 = 72.0%  → RED ZONE\n"
    "F3a Stock Remaining      = 47 / 168 × 100        = 28.0%\n"
    "F4  Order Quantity       = 168 − 47             = 121 units  (Production Plan raised)\n\n"
    "F5  Throughput/Constraint Unit (tie-break):\n"
    "    Selling Price 350, Totally Variable Cost 120, Speed 8 units/min\n"
    "    = (350 − 120) × 8 = 1,840 per constraint minute"
)

h2("4.5 Minimum Order Quantity (MOQ) floor")
para(
    "Every automation that creates a Production Plan or Material Request applies a floor so "
    "tiny orders are never raised. The floor is read per warehouse and converted into the "
    "stock Unit of Measurement before comparison:"
)
code(
    "If the rule's Unit of Measurement is the stock UOM:\n"
    "    Minimum (stock UOM) = min_manufacturing_qty\n"
    "Otherwise:\n"
    "    Minimum (stock UOM) = min_manufacturing_qty × conversion_factor\n\n"
    "Order Quantity used = maximum(shortage, Minimum)"
)

# ===========================================================================
# 5. AUTOMATION ENGINES
# ===========================================================================
h1("5. The Automation Engines (Macro Features)")
para(
    "These are the calculation pipelines that actually create documents. Each one is "
    "described with its purpose, the calculation it performs, and when it is triggered."
)

# --- Buffer calculation
h2("5.1 Buffer calculation engine")
para("Purpose:", bold=True)
para("Compute the buffer status (Inventory Position, Buffer Penetration, Zone, Order Quantity) for every Theory of Constraints (TOC) enabled item at every warehouse rule it has.")
para("Logic:", bold=True)
bullet("Find all items where the TOC master switch is on and the item is not disabled.")
bullet("For each item, resolve its buffer type (Finished Goods, Semi-Finished Goods, Raw Material, Packaging Material) directly or from the Item Group rules.")
bullet("For each enabled warehouse rule, compute F2 (Inventory Position), F3 (Buffer Penetration), the zone, and F4 (Order Quantity).")
bullet("Sort the results by Buffer Penetration descending, then by Throughput per Constraint Unit descending, so the most urgent, most profitable item is first.")
para("Trigger time:", bold=True)
para("On demand whenever a page or report asks for it; also called inside the 07:00, 07:30 and 08:00 scheduled jobs.")

# --- ADU
h2("5.2 Average Daily Usage calculation")
para("Purpose:", bold=True)
para("Keep the consumption rate of every item current from real historical transactions, so buffers reflect what is actually moving.")
para("Logic and calculation:", bold=True)
para(
    "The calculation is universal and does NOT depend on the item's Item Group or buffer "
    "type (Finished, Semi-Finished, Raw, Packaging). Whatever the item is, a single query "
    "reads every outward stock movement from the Stock Ledger Entry — that is, every record "
    "with a negative quantity. One query therefore captures all of: sales shipments "
    "(Delivery Note), production consumption (Stock Entry of type Manufacture or Material "
    "Transfer for Manufacture), manual issues, subcontracting transfers out, and any other "
    "voucher that removes stock."
)
code(
    "Average Daily Usage = total outward quantity in the window  ÷  number of days\n"
    "  total outward = sum of the absolute value of every Stock Ledger Entry\n"
    "                  quantity that is negative (stock left), not cancelled.\n"
    "Window = the warehouse rule's look-back days (default 90).\n"
    "No branching by Item Group or buffer type — the same query for every item.\n"
    "A warehouse row set to manual (auto-rate off) is skipped, keeping its value."
)
para(
    "ADU lives only in the per-warehouse “Minimum Manufacture / Purchase Qty per Warehouse” "
    "table (one row per item and warehouse). A single writer refreshes every row and also "
    "recomputes that row's Maximum Level = Average Daily Usage × Lead Time × Safety Factor, "
    "but only when the row is set to automatic and the item has at least one full look-back "
    "window of history (otherwise the row is left untouched as “warming up”, to avoid "
    "understating the rate). The old standalone item-level ADU fields and their separate "
    "cron were removed on 2026-06-02 — they duplicated this table.", after=4
)
para(
    "Consideration: because the rule is literally “all outward”, the query also includes "
    "negative Stock Reconciliation adjustments and inter-warehouse Material Transfer out-legs "
    "(corrections and relocations, not true demand). This matches the stated requirement; if a "
    "pure demand-only rate is wanted later, those two voucher types can be excluded.",
    italic=True, after=4
)
para("Trigger time:", bold=True)
para("01:00 daily (single per-warehouse ADU + Maximum Level refresh).")

# --- MR generation
h2("5.3 Material Request generation")
para("Purpose:", bold=True)
para("Automatically raise Draft Material Requests for every item whose buffer has fallen into an actionable zone.")
para("Logic:", bold=True)
bullet("Skip everything if the master automatic-generation switch in TOC Settings is off.")
bullet("Take the sorted buffer list and keep only items in the configured zones (default Red, Black and Yellow) that have a positive Order Quantity.")
bullet("Skip an item if a non-terminal Material Request already exists for it at that warehouse (deduplication). Terminal statuses such as Received, Issued, Transferred and Manufactured do NOT block a fresh request — that cycle is complete.")
bullet("Create one Material Request per item, of type Manufacture (manufactured items) or Purchase (purchased items). For purchased items the quantity is divided by the conversion factor so the supplier sees friendly units (for example Kilogram instead of Gram).")
bullet("Stamp the buffer snapshot onto the Material Request (recorded by System, zone, Buffer Penetration, Target Buffer, Inventory Position) for the audit trail, and write a TOC Buffer Log row.")
bullet("Material Requests are left as Draft, never auto-submitted — a planner reviews and submits them.")
bullet("If notify-on-red is on, an email with a colour-coded table is sent to the configured alert roles.")
para("Trigger time:", bold=True)
para("07:00 daily; also on demand from the “Generate Material Requests Now” button on the Production Priority Board.")

# --- component MR
h2("5.4 Component-shortage Material Requests (after Work Orders)")
para("Purpose:", bold=True)
para("After a Production Plan creates Work Orders, raise Purchase Material Requests for any Bill of Materials components that are short.")
para("Logic:", bold=True)
bullet("Runs as the final step inside Production Plan submission, only when the TOC Settings toggle is on.")
bullet("Aggregate the net required quantity for every component across all Work Orders and all Bill of Materials levels.")
bullet("Keep only components flagged as automatically purchased, compare against on-hand stock, apply the Minimum Order Quantity floor, deduplicate against open Purchase Material Requests, and create one request per item and warehouse.")
para("Trigger time:", bold=True)
para("Immediately after Work Orders are created from any automation-driven Production Plan.")

# --- DBM
h2("5.5 Dynamic Buffer Management (weekly auto-resize)")
para("Purpose:", bold=True)
para("Let buffers learn. Grow them quickly when stock keeps hitting the urgent zone; shrink them slowly when they sit comfortable for a long time.")
para("Calculation and logic:", bold=True)
para("Too Much Red (TMR) — the fast-increase trigger:")
code(
    "Look at the last one Replenishment Lead Time of daily buffer logs.\n"
    "Red threshold = Replenishment Lead Time × (configured percent, default 20%).\n"
    "If days in Red/Black exceed the threshold:\n"
    "    new Target Buffer = round(Target × 1.33)   [increase 33%]\n"
    "Safeguard: after 3 consecutive increases, stop and log for manual review."
)
para("Too Much Green (TMG) — the slow-decrease trigger:")
code(
    "Look at the last (Replenishment Lead Time × cycles, default 3) days.\n"
    "If EVERY day in that window was Green:\n"
    "    new Target Buffer = max(floor, round(Target × 0.67))  [decrease 33%]\n"
    "The floor (default 50) prevents the buffer collapsing to near zero."
)
para(
    "The asymmetry is deliberate: increase after roughly one or two red days, but decrease "
    "only after about three full lead-time cycles all green. This avoids cutting a buffer "
    "too early and causing a stockout.", italic=True
)
para("Trigger time:", bold=True)
para("09:00 every Sunday. It runs one hour after the 08:00 snapshot so it never reads data the snapshot has not yet committed.")

# --- Sales projection automation
h2("5.6 Sales Projection automation (Calculation A and Calculation B)")
para("Purpose:", bold=True)
para("Turn a submitted monthly Sales Projection (forecast) into the right replenishment document for items that will fall short — a Production Plan for manufactured items, or a Material Request for purchased items.")
para("Replenishment-mode gate (checked before any document is created, 2026-06-03):", bold=True)
bullet("The item must have a replenishment mode on its master record — “Auto Manufacturing TOC” or “Auto Purchase TOC”. If neither is set, the item is skipped and logged as “No Replenishment Mode” (monitor-only).")
bullet("The item must have a Minimum Purchase/Production Quantity for that warehouse (Item TOC Setting tab, section 6). If it is zero or unset, the item is skipped and logged as “Min Qty Not Set”.")
bullet("The mode decides the document type: Manufacturing → Production Plan + Work Orders; Purchase → Material Request. All purchased shortages in one run are pooled into a single consolidated Material Request.")
para("Calculation A — forecast shortage (per projection row, per item and warehouse):")
code(
    "Shortage A = (Sales Projection + previous-month pending Sales Orders)\n"
    "             − (all current-month Sales Orders + pending Work Orders + current stock)\n"
    "If Shortage A > 0 and the gate passes:\n"
    "    Production quantity = maximum(Shortage A, Minimum Order Quantity)\n"
    "    Manufacturing item → create + submit a Production Plan ([Calc A]) and Work Orders\n"
    "    Purchase item      → add to the run's consolidated Material Request ([Calc A/B][Purchase])"
)
para("Calculation B — Sales Order safety net (runs after Calculation A is committed):")
code(
    "Re-read pending Work Orders and current stock (so Calc A's new Work Order is visible).\n"
    "Shortage B = (previous-month + current-month pending Sales Orders)\n"
    "             − (current stock + pending Work Orders)\n"
    "If Shortage B > 0:\n"
    "    Production quantity = maximum(Shortage B, Minimum Order Quantity)\n"
    "    Create a Production Plan (reason marked [Calc B])"
)
para(
    "The commit between the two calculations is essential: without it Calculation B would "
    "not see the Work Order that Calculation A just created and would double-order.", italic=True
)
para("Trigger time:", bold=True)
para("02:00 daily for all submitted projections of the current month; also on demand from the “Run Production Plan Automation” button on a current-month Sales Projection.")

# --- Calc SO
h2("5.7 Sales Order shortage cover (Calculation SO)")
para("Purpose:", bold=True)
para("Independent of any forecast, scan every pending Sales Order in the company and cover real shortages.")
para("Calculation:", bold=True)
code(
    "For each (item × warehouse) pair that has pending Sales Order quantity:\n"
    "    pending = sum of (stock quantity − delivered × conversion factor)\n"
    "    shortage = pending − current stock − pending Work Orders − pending Purchase Orders\n"
    "    if shortage > 0 and the replenishment-mode + Min-Qty gate passes:\n"
    "        quantity = maximum(shortage, Minimum Order Quantity)\n"
    "        Purchase item    → create a Purchase Material Request\n"
    "        Manufacture item → create a Production Plan + Work Orders"
)
para(
    "The voucher type is decided by the item-master mode (“Auto Purchase TOC” / "
    "“Auto Manufacturing TOC”), not by a per-warehouse setting. Items with no mode, "
    "or with no per-warehouse Minimum Quantity, are skipped and logged with the exact "
    "reason. Supply now includes pending Purchase Orders (added 2026-06-03).", italic=True
)
para("Trigger time:", bold=True)
para("Automatically every day at 07:00 (scheduled, since 2026-06-04), and on demand from the “Run Sales Order Shortage Now” button in TOC Settings. It runs just after the 07:00 buffer Material Request run and uses its own audit run log.")

# --- Calc Action
h2("5.8 Shortage Action automation (Calculation Action)")
para("Purpose:", bold=True)
para("Per warehouse rule, automatically monitor an item in two modes and act when needed.")
para("Logic:", bold=True)
para("Iterates the Item Minimum Manufacture rows where automatic-on-shortage or automatic-on-max-level is on. Two modes, shortage first:")
code(
    "Mode 1 — Shortage:\n"
    "    supply = stock + pending Work Order output + pending Purchase Orders\n"
    "    demand = pending Sales Orders + Work Order component requirement\n"
    "    if demand − supply > 0 → act for that gap (floored by Minimum Order Quantity)\n\n"
    "Mode 2 — Maximum Level (only if Mode 1 did not fire):\n"
    "    cover = (stock + pending Work Orders + pending Purchase Orders)\n"
    "            − (pending Sales Orders + Work Order component requirement)\n"
    "    cover percent = cover ÷ Maximum Level × 100\n"
    "    if cover percent < the row threshold → top up to Maximum Level"
)
para(
    "When a mode fires, the document type follows the item-master replenishment "
    "mode — Manufacturing → Production Plan + Work Orders, Purchase → Material Request. "
    "Items with no mode set are skipped and logged as “No Replenishment Mode”. "
    "Mode 1 supply includes pending Purchase Orders (added 2026-06-03).", italic=True
)
para("Trigger time:", bold=True)
para("On demand from the “Run Shortage Action Now” button in TOC Settings.")

# --- realtime alerts
h2("5.9 Real-time buffer alerts")
para("Purpose:", bold=True)
para("Pop a browser alert the moment a stock movement pushes a Theory of Constraints item into the Red or Black zone.")
para("Logic and trigger:", bold=True)
bullet("When a Stock Ledger Entry is inserted, or a Sales Order / Work Order / Purchase Order is submitted or cancelled, a background job is queued after the database commit.")
bullet("The job recomputes the affected item's buffer and, if it is Red or Black, publishes a live event to the user's browser, which shows a 10-second alert.")
para(
    "The job is queued after commit so it reads the post-transaction stock, never the stale "
    "pre-transaction value.", italic=True
)

# --- replenishment gate + logging
h2("5.10 Replenishment-mode gate and universal automation logging (2026-06-03)")
para("Purpose:", bold=True)
para(
    "Two app-wide guarantees added across every automation that can create a document: "
    "(1) nothing is created unless the item is explicitly configured for it, and "
    "(2) every automation run is auditable with a per-item reason for what it did or why it skipped."
)
para("The replenishment-mode gate (applied before any voucher is created):", bold=True)
bullet("Read the item-master flags “Auto Manufacturing TOC” and “Auto Purchase TOC”. These are the single source of truth for the document type.")
bullet("Manufacturing → Production Plan + Work Orders. Purchase → Material Request. Neither → skip and log “No Replenishment Mode”.")
bullet("The per-warehouse Minimum Purchase/Production Quantity must be greater than zero, else skip and log “Min Qty Not Set”.")
bullet("Every comparison is item-and-warehouse specific. This gate is enforced in the Sales Projection automation (Calc A/B), the Sales Order shortage cover (Calc SO), the Shortage Action automation (Calc Action), and the daily buffer Material Request generator.")
para("Universal logging:", bold=True)
bullet("One audit record per run in “TOC Production Plan Run Log”. Voucher-creating runs add one child “Run Item” row per item with the status (Created / Created as consolidated Material Request / one of the Skipped reasons / Error) and a full plain-text reason and formula breakdown.")
bullet("Created documents are linked from the row: Production Plans in the Production Plan field, Material Requests in the Material Request field.")
bullet("Monitoring jobs that create no documents (per-warehouse Average Daily Usage refresh, procurement scan, buffer snapshot, weekly Dynamic Buffer Management, minimum-order sync) each write one header-only run record summarising what they did — so literally every scheduled job leaves an audit trail.")
para("Consolidated purchase:", bold=True)
para(
    "In the Sales Projection automation, all purchased-item shortages found in a single "
    "run are pooled into ONE Material Request (one line per item, in each item's purchase "
    "unit of measurement), rather than many separate requests — easier for the buyer to action.")

# ===========================================================================
# 6. SCHEDULED TASKS
# ===========================================================================
h1("6. Scheduled Background Jobs (Trigger Times)")
para(
    "All times are in the server timezone. Each job runs in the long-queue worker as the "
    "site's system user. The one-hour gap before Dynamic Buffer Management is intentional."
)
para(
    "Every job below also writes one audit record to “TOC Production Plan Run Log” "
    "(2026-06-03): voucher-creating jobs with per-item rows and reasons, monitoring "
    "jobs as a one-line summary.", italic=True
)
table(
    ["Time", "Job", "Purpose", "What it writes"],
    [
        ("00:00 daily", "Minimum order quantity sync", "Refresh the per-item minimum order quantity.", "Item.custom_min_order_qty; run-log summary"),
        ("01:00 daily", "Per-warehouse Average Daily Usage + Maximum Level", "Refresh each (item, warehouse) consumption rate from all outward stock movement (item-group independent) and recompute Maximum Level. Sole ADU job.", "Item Minimum Manufacture rows; run-log summary"),
        ("02:00 daily", "Sales Projection automation (Calc A + Calc B)", "Create Production Plans (manufacture) or a consolidated Material Request (purchase) for forecast/Sales-Order shortfalls, after the replenishment-mode gate.", "Production Plan, Work Order, Material Request, run log (per-item)"),
        ("07:00 daily", "Production / buffer Material Request run", "Raise Material Requests / Production Plans for actionable buffers (mode-gated).", "Material Request / Production Plan (Draft), TOC Buffer Log, run log (per-item), email alerts"),
        ("07:00 daily", "Sales Order shortage cover (Calc SO)", "Cover real shortages against every pending Sales Order — Production Plan + Work Orders (manufacture) or Material Request (purchase), after the replenishment-mode + Min-Qty gate. Auto since 2026-06-04 (previously opt-in); runs just after the buffer run.", "Production Plan, Work Order, Material Request, run log (per-item)"),
        ("07:30 daily", "Procurement monitoring", "Log purchase items in Red/Black for the purchase team (no documents created).", "Run-log summary; logger"),
        ("08:00 daily", "Buffer snapshot", "Archive every buffer state for history and trend analysis.", "TOC Buffer Log (one row per item × warehouse); run-log summary"),
        ("09:00 Sunday", "Weekly Dynamic Buffer Management", "Resize buffers (Too Much Red / Too Much Green).", "TOC Item Buffer target buffers + counters; run-log summary"),
    ],
    widths=[1.05, 1.85, 2.2, 1.9],
    font=8.5,
)
para(
    "Note: the per-warehouse Average Daily Usage value lives on the Item Minimum "
    "Manufacture row; use the recalculate-item-buffers "
    "endpoint to synchronise them when needed.", italic=True
)

h2("6.1 Configurable Triggers & Per-Trigger Pending Statuses (2026-06-04)")
para("Purpose:", bold=True)
para(
    "Every automation engine is listed on the TOC Settings page (the "
    "“Automation Engines & Triggers” section) with its schedule, an "
    "Enabled switch, and a manual Run Now button. Editing a trigger’s time "
    "rewrites Frappe’s native Scheduled Job Type immediately — no restart "
    "or migrate. All nine engines are auto-listed in the Trigger Configurations "
    "table after the app is installed (they cannot be added or removed by hand; "
    "the rows are system-managed)."
)
para("Per-trigger pending statuses:", bold=True)
para(
    "Each engine row can override which Sales Order, Work Order and Purchase "
    "Order statuses count as pending, using the same Status:Workflow-State "
    "multiselect as the global fields (open a row with the pencil to edit). "
    "Resolution order is: the trigger row’s override (if filled), else the "
    "global TOC Settings field, else a built-in default. The global fields "
    "remain the single source of truth for the reports (Work Order Kitting "
    "Planner, Production Overview). Only the three Sales-Order / shortage "
    "engines (Calc A+B, Calc SO, Calc Action) read these lists; the buffer and "
    "procurement runs use live Bin quantities and are marked Not Applicable."
)
para("Manual run — access:", bold=True)
para(
    "Run Now enqueues the selected engine on the long queue; the button and the "
    "scheduled job call exactly the same code. Manual runs are restricted to "
    "System Managers. Shortage Action (Calc Action) is schedulable too, seeded "
    "disabled so it only runs once a System Manager enables its row."
)
para("Purchase Material Request netting (2026-06-04):", bold=True)
para(
    "For a purchase-mode item the engine creates a Material Request (a Purchase "
    "Order cannot be raised automatically because the supplier is the buyer's "
    "choice). To avoid re-requesting the same shortage forever, the three "
    "purchase engines (Sales Projection, Sales Order Shortage, Shortage Action) "
    "subtract the quantity already pending on open/draft Purchase Material "
    "Requests from the shortage. A fourth pending list, 'Pending Purchase "
    "Material Request Statuses', is configured globally in TOC Settings and can "
    "be overridden per trigger; it must include Draft because the engine leaves "
    "its own requests as Draft. The remaining quantity is measured as ordered "
    "minus received so a request that has become a Purchase Order is not counted "
    "twice."
)
para("Average Daily Usage — Stock Reconciliation excluded (2026-06-04):", bold=True)
para(
    "The Average Daily Usage refresh no longer counts Stock Reconciliation "
    "outward legs (inventory corrections), only true consumption (Delivery "
    "Notes, Work Order / Stock Entry consumption, issues, transfers), so a "
    "physical-count adjustment does not inflate the buffer."
)
para("Created-by metadata (2026-06-04):", bold=True)
para(
    "Material Requests, Production Plans and Work Orders carry a read-only "
    "'Recorded By' (User / System) and a formatted 'Creation Reason' showing, in "
    "a small table, exactly why the document was created and the live figures the "
    "engine used. System documents are read-only; user-created ones can be edited "
    "before submit."
)

# ===========================================================================
# 7. DOCTYPES
# ===========================================================================
h1("7. Data Records (DocTypes)")

# TOC Settings
h2("7.1 TOC Settings (single configuration record)")
para("Purpose:", bold=True)
para("One record per site holding every global setting. It is the single source of truth for what counts as pending across all reports and jobs.")
para("Key fields, grouped:", bold=True)
table(
    ["Group", "Field (purpose)"],
    [
        ("Zone thresholds", "Red zone threshold (default 67), Yellow zone threshold (default 33). Red must be greater than Yellow."),
        ("Material Request generation", "Master automatic-generation switch; which zones trigger requests; notify-on-red; the alert roles."),
        ("Dynamic Buffer Management", "Enable switch; Too-Much-Red percent of lead time (default 20); Too-Much-Green cycles (default 3); adjustment percent (default 33); maximum consecutive increases (default 3); minimum buffer floor (default 50)."),
        ("Calculation defaults", "Default Variability Factor (1.5); Average Daily Usage look-back days (90); global Demand Adjustment Factor (1.0) and its event label."),
        ("Pending status lists", "Pending Sales Order / Work Order / Purchase Order statuses and workflow states — read by every report and job."),
        ("Warehouse classification", "Child table tagging each warehouse as Inventory, Work-In-Process, or Excluded (so scrap and expiry stock never inflate a buffer)."),
        ("Action buttons", "Run Sales Order Shortage Now; Run Shortage Action Now; demo data manifest (hidden)."),
    ],
    widths=[1.7, 5.3],
)
para("Validation logic: Red threshold must exceed Yellow; Variability Factor must be at least 1.0; adjustment percent must be 1 to 100; warehouses cannot repeat in the classification table.")

# TOC Item Buffer
h2("7.2 TOC Item Buffer (per-warehouse buffer rule, child of Item)")
para("Purpose:", bold=True)
para("One row equals one warehouse buffer rule for a Theory of Constraints item. This is where Average Daily Usage, Replenishment Lead Time, Variability Factor and Demand Adjustment Factor are stored and where F1 and F6 are calculated on save.")
table(
    ["Field", "Type", "Meaning / calculation"],
    [
        ("Warehouse", "Link", "Home warehouse for this rule."),
        ("Average Daily Usage", "Number", "Units consumed per day."),
        ("Replenishment Lead Time", "Number", "Days from order to availability."),
        ("Variability Factor", "Number", "Safety multiplier (default 1.5)."),
        ("Target Buffer", "Number (auto)", "F1 = Average Daily Usage × Lead Time × Variability Factor."),
        ("Demand Adjustment Factor", "Number", "Seasonal multiplier (default 1.0)."),
        ("Demand-Adjusted Buffer", "Number (auto)", "F6 = Target × Demand Adjustment Factor; stored as 0 when factor is 1.0 (meaning “use the base target”)."),
        ("Red / Yellow zone quantity", "Number (auto)", "Stock levels below which the rule is Red / Yellow, for operators who prefer units to percentages."),
        ("Enabled", "Yes/No", "Uncheck to disable a rule without deleting it."),
    ],
    widths=[2.0, 1.2, 3.8],
)

# TOC Buffer Log
h2("7.3 TOC Buffer Log (daily snapshot archive)")
para("Purpose:", bold=True)
para("One row equals one point-in-time reading of an item's buffer at a warehouse on a date. This historical record feeds Dynamic Buffer Management, the trend reports and the workspace number cards. Records are created automatically, never edited by hand.")
para("Each row stores the item, warehouse, date, buffer type, the five Inventory Position components, the Target Buffer, Buffer Penetration, Stock Remaining, the zone, the suggested Order Quantity and (when applicable) the linked Material Request.")
para("Created by: the 08:00 snapshot (all items) and the 07:00 Material Request run (only items that got a request). No automatic cleanup exists — the table grows roughly one to two rows per item per warehouse per day.")

# Sales Projection
h2("7.4 Sales Projection (and its child tables)")
para("Purpose:", bold=True)
para("Captures the minimum production target per item for a calendar month, year and warehouse. It is unique on month + year + warehouse and is the input for the 02:00 Production Plan automation.")
table(
    ["Element", "Purpose"],
    [
        ("Header: month, year, source warehouse", "Identify the projection; the warehouse is also the target for the auto Production Plans."),
        ("Child table: Sales Projected Items", "The projected quantity per item (with Unit of Measurement conversion to stock units)."),
        ("Child table: SP Minimum Manufacture", "Optional per-item, per-warehouse minimum batch size (the Minimum Order Quantity floor)."),
        ("Last automatic run", "Timestamp of the last automation pass."),
    ],
    widths=[2.6, 4.4],
)
para("Validation logic: required header fields; no duplicate items in the child table; no second active (Draft or Submitted) projection for the same month + year + warehouse. Cancelled projections are excluded from the duplicate check, so cancel-then-recreate and cancel-then-amend both work.")
para(
    "A dedicated role, Sales Projection Administrator, can cancel, amend, edit and resubmit a "
    "projection. A deliberate asymmetric fix resolves the cancel deadlock between a projection "
    "and its Production Plans: the projection side fully skips the inbound-link scan, while the "
    "Production Plan side only clears the specific back-links, keeping the Work Order and "
    "Material Request guards intact.", italic=True
)

# Item Min Order Qty
h2("7.5 Item Minimum Order Quantity (child of Item)")
para("Purpose:", bold=True)
para("Per-warehouse minimum order/production batch. The user enters warehouse, Unit of Measurement and minimum quantity; the stock Unit of Measurement, conversion factor and stock-unit quantity are computed automatically. This floor is applied to both buffer and component purchase requests.")

# ===========================================================================
# 8. CUSTOM FIELDS
# ===========================================================================
h1("8. Custom Fields Added to ERPNext DocTypes")
para(
    "The application adds custom fields (shipped as fixtures, module “Chaizup Toc”) to four "
    "ERPNext doctypes. Buffer-snapshot fields are written by the system at document creation "
    "time so the exact triggering state is preserved for audit."
)

h2("8.1 Item — the TOC Setting tab (five sections)")
table(
    ["Section", "Fields (purpose)"],
    [
        ("1. Enable & classify", "Master switch; buffer type (Finished / Semi-Finished / Raw / Packaging); automatic-purchase and automatic-manufacture flags (mutually exclusive)."),
        ("2. Average Daily Usage", "Manual-rate switch; look-back period (30/90/180/365 days); the rate value; last-updated timestamp."),
        ("3. Throughput per Constraint Unit", "Selling price; Totally Variable Cost; constraint speed; the computed Throughput per Constraint Unit (Finished Goods only)."),
        ("4. Bill of Materials & dependency", "Default Bill of Materials link (validated to belong to this item); enable the multi-level component-availability check."),
        ("5. Buffer rules", "The per-warehouse buffer rule table; and the minimum order quantity table."),
    ],
    widths=[1.9, 5.1],
)

h2("8.2 Material Request, Work Order, Purchase Order — buffer snapshot fields")
para("Each of these carries the buffer state captured when the document was created:")
bullet("Recorded by (By System or By User) — marks whether Theory of Constraints automation created it.")
bullet("Zone, Buffer Penetration, Target Buffer, Inventory Position, Stock Remaining — the snapshot used for the audit trail and for sequencing urgent Work Orders on the constraint machine.")
para("When a planner converts an automation-created Material Request into a Purchase Order, these fields are copied to the Purchase Order header automatically, so every order can be traced to the buffer signal that caused it.")

# ===========================================================================
# 9. OVERRIDES
# ===========================================================================
h1("9. ERPNext Behaviour Overrides")
para("These hook into existing ERPNext doctypes; no new doctypes are introduced here.")
table(
    ["Override", "What it does", "When"],
    [
        ("Item validation", "Validate the manual rate, enforce purchase-or-manufacture exclusivity, compute Throughput per Constraint Unit, validate the Bill of Materials link, warn if no buffer rules exist.", "Every Item save where Theory of Constraints is enabled."),
        ("Material Request validation", "Warn (never block) when a user manually raises a request for a Theory of Constraints managed item.", "Every Material Request save not marked By System."),
        ("Production Plan class", "Stamp the Theory of Constraints Unit of Measurement display fields after the “Get Sub Assembly Items” and “Create Work Order” buttons, which otherwise skip the normal save hooks.", "On those button clicks."),
        ("Purchase Order before insert", "Copy the buffer snapshot fields from the source automation-created Material Request to the Purchase Order header.", "When a Purchase Order is created from such a request."),
        ("Reorder override", "Replace ERPNext's default reorder routine entirely; if someone re-enables ERPNext auto-indent, it logs a warning and switches it back off.", "Whenever ERPNext would run its reorder."),
    ],
    widths=[1.7, 3.9, 1.4],
    font=8.5,
)

# ===========================================================================
# 10. PAGES
# ===========================================================================
h1("10. Interactive Pages (Macro Features)")
para(
    "Eleven Desk Pages provide the operational interface. All follow the same critical rule: "
    "no raw apostrophe or single quote may appear anywhere in the page HTML, because Frappe "
    "wraps the template in a single-quoted JavaScript string and one stray quote blanks the "
    "whole page. After any HTML change the Redis cache must be flushed."
)

def page_block(title, route, purpose, features, calc=None, trigger=None):
    h2(title)
    para("Route: " + route, italic=True, color=SLATE_500, size=9.5, after=3)
    para("Purpose:", bold=True, after=2)
    para(purpose, after=4)
    if calc:
        para("Calculation / logic:", bold=True, after=2)
        for c in calc:
            bullet(c)
    para("Key features:", bold=True, after=2)
    for f in features:
        bullet(f)
    if trigger:
        para("Trigger / refresh:", bold=True, after=2)
        para(trigger, after=4)


page_block(
    "10.1 Work Order Kitting Planner",
    "/app/wo-kitting-planner",
    "Simulate material availability across every open Work Order, find shortages, and decide what can actually be produced. The planning console for the production manager.",
    [
        "Seven tabs: Production Plan, Material Shortage, Emergency, Dispatch, Artificial-Intelligence Advisor, Item View, Purchase Priority.",
        "Two supply modes: current stock only, or current plus expected (Purchase Orders, Material Requests, Work Orders).",
        "Two calculation modes: isolated (each Work Order checked against full stock) and sequential (Work Orders consume stock in priority order).",
        "Dispatch bottleneck analysis comparing Finished Goods stock and plan against Sales Order demand.",
        "Cost audit comparing the standard Bill of Materials cost against the actual consumed cost.",
        "Artificial-Intelligence advisor that summarises the top issues and answers production questions.",
        "Dual Unit-of-Measurement display on every quantity (for example 1,000 Gram shown above 1.00 Kilogram).",
    ],
    calc=[
        "Work Order finished item is read from the production-item field; Purchase Order price from the rate field; Stock Entry link from the work-order field.",
        "Item groups in the shortage tab come from the actual components in the simulation, never inherited from the parent Work Order.",
        "All pending Sales Order / Work Order / Purchase Order definitions are read from TOC Settings (no status pickers on the page; a read-only banner shows the active filter).",
    ],
    trigger="On page load and whenever the user presses the recalculation control.",
)

page_block(
    "10.2 Item Short / Surplus",
    "/app/item-short-surplus",
    "For every item with stock or pending vouchers, answer one question: will today's supply (current stock + pending Work Order output + pending Purchase Order receipts) cover today's demand (pending Sales Orders + remaining Work Order consumption)?",
    [
        "Classifies each item as Shortage or Surplus, with the magnitude in both the stock and a higher Unit of Measurement.",
        "Live multi-select filters for item, item group, warehouse, company, and pending statuses; six on/off toggles for active or no Sales Order / Work Order / Purchase Order.",
        "Per-cell drill-down to the contributing vouchers; sortable, sticky-header, word-wrapped table.",
        "Four-sheet spreadsheet export (Main, Filters & Run Info, Shortage sorted, Surplus sorted).",
    ],
    calc=[
        "Supply = current stock + pending Work Orders + pending Purchase Orders.",
        "Demand = pending Sales Orders + remaining Work Order consumption.",
        "Net = Supply − Demand; Surplus when Net is zero or above, Shortage otherwise; Shortfall = absolute value of a negative Net.",
        "A voucher counts as pending when submitted and in a chosen status, or draft and in a chosen workflow state. Closed, Cancelled and Stopped are always excluded.",
        "Every filter resolves live at query time — there are no cached aggregates or stored mirrors, so the screen always reflects the current moment.",
    ],
    trigger="On load and on every filter change.",
)

page_block(
    "10.3 Item Shortage Dashboard",
    "/app/item-shortage-dashboard",
    "Single pane of glass: one row per item-and-warehouse rule showing stock, maximum level, pending Sales/Work/Purchase Orders, sales projection, dispatch, Average Daily Usage, lead time and safety factor together.",
    [
        "Grid with frozen header, zebra rows, and dual Unit-of-Measurement cells.",
        "Quick-filter chips (Shortage, Open Sales Order, Open Purchase Order, Open Work Order, Below 50% maximum) with live counts.",
        "Every numeric cell is clickable for a drill-down to the contributing documents.",
        "Server-side five-sheet spreadsheet export and an email-snapshot action.",
    ],
    calc=[
        "Maximum level, Average Daily Usage, lead time and safety factor come from the Item Minimum Manufacture child table.",
        "Decision quantity = the larger of the need-to-maximum-level and the total shortage including expected receipts.",
    ],
    trigger="On load and on filter change.",
)

page_block(
    "10.4 Item Projection View",
    "/app/item-projection-view",
    "Per item and warehouse, show current stock plus pipeline inflows minus demand in a higher Unit of Measurement, with every cell explaining its formula and linking to the source vouchers.",
    [
        "Sixteen columns including physical shortage, projected shortage, Work Order remaining production, Purchase Order remaining, will-consume, will-dispatch, net available and days of cover.",
        "Tooltip on every number giving the formula and the actual contributing values.",
        "Grouping by item group with summed headers; eight-sheet branded spreadsheet export.",
    ],
    calc=[
        "Demand = Work Order consumption + Sales Order dispatch.",
        "Will-receive = Purchase Order remaining + Work Order production.",
        "Physical shortage = max(0, demand − stock); Projected shortage = max(0, demand − (stock + will-receive)); Net available = stock + will-receive − demand.",
        "Days of cover = stock ÷ Average Daily Usage (only meaningful at the item-warehouse leaf level).",
    ],
    trigger="On load and on filter change.",
)

page_block(
    "10.5 Production Overview",
    "/app/production-overview",
    "Single-screen production planning console consolidating open Work Orders, Sales Orders, projections, dispatch, shortage status and cost variance into one row per item, across three tabs (Overview, Artificial-Intelligence Advisor, Charts).",
    [
        "Eighteen-column item table with per-item drill-down modals (Work Order chain, sub-assembly tree, batch consumption, cost).",
        "Planned versus actual production, projection-versus-sales ratio, shortage flag, possible production quantity, three-way cost comparison.",
        "Planning mode with an Independent and a Priority-queue sub-mode (drag to re-sequence; the simulation re-runs against the new order).",
        "Quick-filter pills (open Sales Order, open Work Order, in Production Plan, etc.) and an eight-sheet spreadsheet export with live formulas.",
    ],
    calc=[
        "Pending Sales Order demand uses stock quantity minus delivered, clamped at zero.",
        "Target production = the larger of the sales projection and the total current sales.",
        "Possible quantity is a Bill-of-Materials kitting simulation against the available stock pool.",
    ],
    trigger="On pressing Load; charts and the advisor refresh lazily after a load.",
)

page_block(
    "10.6 Supply Chain Tracker (Manufacturing Pipeline)",
    "/app/supply-chain-tracker",
    "A seven-stage pipeline view tracing every Theory of Constraints item from the first replenishment trigger to goods in hand.",
    [
        "Stages: Items, Material Request, Request-for-Quotation or Production Plan, Supplier Quotation or Work Order, Purchase Order or Job Card, Receipt/Quality/Stock Entry, Finished output.",
        "Click any card to highlight its full upstream and downstream chain; a detail panel shows all fields plus the buffer breakdown.",
        "Zone filter is instant (in-browser); the replenishment-mode filter re-fetches.",
        "Summary strip with Red/Yellow/Green counts and average Buffer Penetration.",
    ],
    trigger="On load; replenishment-mode change re-fetches, zone change filters in place.",
)

page_block(
    "10.7 TOC Dashboard",
    "/app/toc-dashboard",
    "Live, auto-refreshing overview of all buffer zones — the morning at-a-glance view.",
    [
        "Simplified zone-scanning table, summary cards, and a live donut chart.",
        "Auto-refreshes every five minutes.",
        "Action buttons gated to authorised roles.",
    ],
    trigger="Every five minutes (automatic).",
)

page_block(
    "10.8 Kitting Report",
    "/app/kitting-report",
    "Full production readiness check: which Finished / Semi-Finished items have pending demand, whether components are available, and the procurement status of each shortage.",
    [
        "Demand-driven (Sales Orders → demand → Bill-of-Materials check), with per-row drill-down.",
        "One-click creation of a consolidated purchase Material Request, or a Work Order, from the shortage list.",
        "Manual refresh only (the Bill-of-Materials walk is expensive).",
    ],
    calc=[
        "Production required = max(0, total pending demand − in-stock); should-produce subtracts what is already produced.",
        "Kit percentage is driven by the most constrained component in the Bill of Materials.",
    ],
    trigger="Manual refresh.",
)

page_block(
    "10.9 TOC Item Settings (bulk configuration)",
    "/app/toc-item-settings",
    "Bulk item configuration. Filter items, then edit all Theory of Constraints fields per item in a modal, or apply some fields in bulk across many items at once.",
    [
        "Modal with the same five sections as the Item form, plus a help panel explaining each formula.",
        "Bulk-apply for automatic-rate, lead time and safety factor (batch-specific fields are deliberately excluded).",
    ],
    trigger="On demand.",
)

page_block(
    "10.10 TOC User Guide",
    "/app/toc-user-guide",
    "A self-contained help page covering every topic: overview, quick start, all formulas with live calculators, buffer types, the zone system, the daily schedule, projection automation, Dynamic Buffer Management, alerts, settings reference, field reference, the trigger map and troubleshooting.",
    [
        "Static page, no server calls; navigation sidebar with scroll tracking and a search filter.",
    ],
    trigger="None (static reference).",
)

# ===========================================================================
# 11. REPORTS
# ===========================================================================
h1("11. Script Reports")
para("Five Script Reports cover the daily operational and analytical surface.")
table(
    ["Report", "Audience", "Data source", "Purpose"],
    [
        ("Production Priority Board", "Production supervisor", "Live buffer calculation", "What to produce today (Finished and Semi-Finished items), ranked by Buffer Penetration; buttons to apply Demand Adjustment Factor and generate Material Requests."),
        ("Procurement Action List", "Procurement officer", "Live buffer calculation (Raw + Packaging)", "What to buy today, ranked by Buffer Penetration."),
        ("Buffer Status Report", "Operations manager", "Historical buffer log", "How buffer health has changed over time (trend)."),
        ("Dynamic Buffer Management Analysis", "Theory of Constraints manager", "Rules + 30 days of buffer log", "Whether Dynamic Buffer Management is correctly resizing buffers."),
        ("Production Indent Subs", "Production planner", "Live", "Per-Work-Order bill of components for pending Work Orders, with both Units of Measurement."),
    ],
    widths=[1.7, 1.4, 1.7, 2.2],
    font=8.5,
)
para("Trigger time: all run on demand when the user opens them and applies filters.")

# ===========================================================================
# 12. APIs
# ===========================================================================
h1("12. Whitelisted Programming Endpoints")
para(
    "All endpoints are callable from the pages, from external clients over the REST "
    "interface, and from server code. Grouped by module:"
)
table(
    ["Module", "Representative endpoints (purpose)"],
    [
        ("Core (toc_api)", "Priority board; single-item buffer; manual Material Request run; recalculate item buffers; apply/reset global Demand Adjustment Factor; buffer summary; Bill-of-Materials check; the four zone number-card counts."),
        ("Kitting (kitting_api)", "Kitting summary; item kitting detail (recursive Bill-of-Materials walk); create consolidated purchase requests; create a Work Order."),
        ("Pipeline (pipeline_api)", "Full supply-chain graph (nodes and edges) and filter options for the tracker."),
        ("Work Order Kitting (wo_kitting_api)", "Open Work Orders; kitting simulation; dispatch bottleneck; item summary; the shared pending-filter banner; Artificial-Intelligence integration."),
        ("Production Overview (production_overview_api)", "Overview, item detail, shortage detail, cost breakup, charts, spreadsheet export, Artificial-Intelligence insight and chat."),
        ("Item shortage / projection / short-surplus", "Dashboard data, drill-down breakdown, filter options, email snapshot, spreadsheet export."),
        ("Permissions", "Application-tile visibility and buffer-log access checks."),
        ("Demo data", "Admin-only seed and delete of a complete test dataset."),
    ],
    widths=[2.0, 5.0],
    font=8.5,
)

# ===========================================================================
# 13. FIXTURES
# ===========================================================================
h1("13. Fixtures (Data Shipped with the App)")
para(
    "Fixtures are doctype records re-imported on every install, migrate and restore. Treat "
    "them as part of the schema, not as runtime data."
)
table(
    ["Fixture file", "Doctype", "Selected by", "What it ships"],
    [
        ("custom_field.json", "Custom Field", "module = Chaizup Toc", "Every custom field on Item, Material Request, Work Order, Purchase Order and Production Plan."),
        ("property_setter.json", "Property Setter", "module = Chaizup Toc", "List-view defaults and field tweaks on Work Order and Bill of Materials."),
        ("list_view_settings.json", "List View Settings", "name in [Work Order, BOM]", "Opinionated default columns for those two list views."),
    ],
    widths=[1.8, 1.4, 1.8, 2.0],
    font=8.5,
)
para(
    "Important: the fixture importer only inserts; it never updates a row that already "
    "exists. Any change to a fixture must be paired with a one-shot patch that rewrites the "
    "affected rows, or live sites keep the old behaviour. Re-export after a user-interface "
    "edit with the export-fixtures command and commit the result.", italic=True
)

# ===========================================================================
# 14. INSTALL / ROLES
# ===========================================================================
h1("14. Installation, Roles and Setup")
para("On install the application performs, in order:")
bullet("Disable ERPNext's built-in automatic reorder (so it never competes with the buffer engine).")
bullet("Create all custom fields.")
bullet("Create the roles (see below).")
bullet("Create four workspace number cards (Red, Yellow, Green zone counts and open Material Requests).")
bullet("Create a zone-distribution donut chart.")
para("On uninstall it re-enables ERPNext's automatic reorder so the standard behaviour returns.")
para("Roles:", bold=True)
table(
    ["Role", "Access"],
    [
        ("Theory of Constraints Manager", "Full: trigger requests, apply Demand Adjustment Factor, view all reports, manage settings."),
        ("Theory of Constraints User", "Read-only: dashboards, reports, buffer logs."),
        ("Sales Projection Administrator", "Full lifecycle on Sales Projection: cancel, amend, edit, resubmit."),
    ],
    widths=[2.5, 4.5],
)
para(
    "Recommended go-live order: install ERPNext, install this app, configure TOC Settings "
    "(warehouse classification first, then thresholds, then Material Request settings), enable "
    "two or three test items, run the daily jobs manually to verify, then enable the rest. "
    "Wait about a month of buffer-log history before enabling Dynamic Buffer Management.", italic=True
)

# ===========================================================================
# 15. HOOKS WIRING
# ===========================================================================
h1("15. How Everything Is Wired (hooks)")
para("The hooks file is the single place that connects the app to Frappe and ERPNext:")
table(
    ["Hook", "Wires"],
    [
        ("Scheduler events", "The daily and weekly background jobs (see section 6)."),
        ("Document events", "Stock Ledger Entry insert, Sales/Work/Purchase Order submit and cancel → real-time alerts; Item save → validation; Material Request save → compliance warning; Production Plan before-cancel → clear projection back-links."),
        ("Override whitelisted methods", "Replace ERPNext's reorder routine with the Theory of Constraints version."),
        ("Override doctype class", "Subclass the Production Plan to stamp Unit-of-Measurement display fields on button-driven flows."),
        ("Client scripts", "A global branding script (zone colours, real-time alerts, keyboard shortcut) plus per-form scripts on Item, Material Request and Stock Entry."),
        ("Apps screen", "The Chaizup TOC tile on the home screen, gated by role."),
        ("Fixtures", "Export custom fields and property setters tagged to this module."),
    ],
    widths=[1.9, 5.1],
    font=8.5,
)

# ===========================================================================
# 16. CRITICAL RULES
# ===========================================================================
h1("16. Critical Rules and Safeguards")
para("These are the non-negotiable constraints that keep the application correct and safe.")
bullet("Page HTML must contain zero raw apostrophes or single quotes; one stray quote blanks the page. Flush the Redis cache after any HTML change.", "Page templates — ")
bullet("Always use the stock-quantity columns (not the transaction-quantity columns) for stock maths, and multiply delivered quantities by the conversion factor before comparing.", "Units of Measurement — ")
bullet("Pending Sales / Work / Purchase Order status definitions live only in TOC Settings; no report or job may hardcode them.", "Single source of truth — ")
bullet("Filters always reflect the current state. No caches, no stored mirrors, no copied fields backing a filter — only live queries.", "Filter accuracy — ")
bullet("Workflow-state branches in the eligibility queries are guarded by a column-existence check, so sites without a workflow do not error.", "Workflow safety — ")
bullet("Every automation deduplicates against non-terminal documents; terminal statuses (completed cycles) never block a fresh document.", "Deduplication — ")
bullet("Daily-job writes do not bump the document's modified timestamp, keeping the audit history clean.", "Audit hygiene — ")
bullet("Never edit ERPNext or Frappe core; all behaviour is added via custom fields, property setters, hooks and overrides so future upgrades stay clean.", "No core edits — ")

spacer(6)
para(
    "End of reference. This document was generated from the in-repository developer "
    "documentation (the per-folder Markdown files, the formula reference and the application "
    "root reference). For the deepest implementation detail, consult those files and the "
    "source code they cite.", italic=True, color=SLATE_500, size=9.5
)

# ---------------------------------------------------------------------------
out_path = "/workspace/development/frappe-bench/apps/chaizup_toc/documentation/Chaizup_TOC_Feature_Reference.docx"
doc.save(out_path)
print("Saved:", out_path)
print("Sections:", sum(1 for p in doc.paragraphs if p.style.name == "Heading 1"))
print("Paragraphs:", len(doc.paragraphs))
print("Tables:", len(doc.tables))
